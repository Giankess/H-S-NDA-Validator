from typing import List, Dict, Any, Tuple
import torch
from torch.utils.data import Dataset, DataLoader
from transformers import (
    AutoTokenizer,
    AutoModelForSequenceClassification,
    AutoModelForTokenClassification,
    TrainingArguments,
    Trainer,
    DataCollatorForTokenClassification
)
from sentence_transformers import SentenceTransformer, InputExample, losses
from sentence_transformers.readers import InputExample
import numpy as np
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.oxml.shared import qn
from ..core.config import settings
from ..services.document_storage import DocumentStorage
from ..services.vector_storage import VectorStorage

class TrainingService:
    def __init__(self):
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.document_storage = DocumentStorage()
        self.vector_storage = VectorStorage()
        
        # Initialize models
        self.classifier_tokenizer = AutoTokenizer.from_pretrained("nlpaueb/legal-bert-base-uncased")
        self.classifier_model = AutoModelForSequenceClassification.from_pretrained(
            "nlpaueb/legal-bert-base-uncased",
            num_labels=3  # [keep, modify, remove]
        ).to(self.device)
        
        self.ner_tokenizer = AutoTokenizer.from_pretrained("nlpaueb/legal-bert-base-uncased")
        self.ner_model = AutoModelForTokenClassification.from_pretrained(
            "nlpaueb/legal-bert-base-uncased",
            num_labels=5  # [O, B-CLAUSE, I-CLAUSE, B-SECTION, I-SECTION]
        ).to(self.device)
        
        self.sentence_transformer = SentenceTransformer('all-MiniLM-L6-v2').to(self.device)

    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extract text from a DOCX file"""
        doc = Document(docx_content)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def extract_changes_from_redline(self, docx_content: bytes) -> List[Dict[str, str]]:
        """Extract changes from a redline DOCX file"""
        doc = Document(docx_content)
        changes = []
        
        for paragraph in doc.paragraphs:
            original_text = []
            redline_text = []
            
            for run in paragraph.runs:
                # Check if this run has tracked changes
                if run._element.xpath('.//w:ins') or run._element.xpath('.//w:del'):
                    # This is a changed run
                    if run._element.xpath('.//w:ins'):
                        # This is an insertion
                        redline_text.append(run.text)
                    if run._element.xpath('.//w:del'):
                        # This is a deletion
                        original_text.append(run.text)
                else:
                    # This is unchanged text
                    original_text.append(run.text)
                    redline_text.append(run.text)
            
            if original_text or redline_text:
                changes.append({
                    "original": "".join(original_text),
                    "redline": "".join(redline_text)
                })
        
        return changes

    def prepare_classification_dataset(self, original_texts: List[str], redline_texts: List[str]) -> Dataset:
        """Prepare dataset for clause classification training"""
        class ClassificationDataset(Dataset):
            def __init__(self, texts, labels, tokenizer):
                self.encodings = tokenizer(texts, truncation=True, padding=True)
                self.labels = labels

            def __getitem__(self, idx):
                item = {key: torch.tensor(val[idx]) for key, val in self.encodings.items()}
                item['labels'] = torch.tensor(self.labels[idx])
                return item

            def __len__(self):
                return len(self.labels)

        # Determine labels based on text differences
        labels = []
        for orig, redline in zip(original_texts, redline_texts):
            if orig == redline:
                labels.append(0)  # keep
            elif not redline:
                labels.append(2)  # remove
            else:
                labels.append(1)  # modify

        return ClassificationDataset(original_texts, labels, self.classifier_tokenizer)

    def train_classifier(self, train_dataset: Dataset, eval_dataset: Dataset = None):
        """Train the clause classification model"""
        training_args = TrainingArguments(
            output_dir="./models/classifier",
            num_train_epochs=3,
            per_device_train_batch_size=8,
            per_device_eval_batch_size=8,
            warmup_steps=500,
            weight_decay=0.01,
            logging_dir="./logs",
            logging_steps=10,
            evaluation_strategy="epoch" if eval_dataset else "no",
        )

        trainer = Trainer(
            model=self.classifier_model,
            args=training_args,
            train_dataset=train_dataset,
            eval_dataset=eval_dataset,
        )

        trainer.train()
        trainer.save_model()

    def prepare_ner_dataset(self, texts: List[str], labels: List[List[int]]) -> Dataset:
        """Prepare dataset for NER training"""
        class NERDataset(Dataset):
            def __init__(self, texts, labels, tokenizer):
                self.encodings = tokenizer(texts, truncation=True, padding=True)
                self.labels = labels

            def __getitem__(self, idx):
                item = {key: torch.tensor(val[idx]) for key, val in self.encodings.items()}
                item['labels'] = torch.tensor(self.labels[idx])
                return item

            def __len__(self):
                return len(self.labels)

        return NERDataset(texts, labels, self.ner_tokenizer)

    def train_ner(self, train_dataset: Dataset, eval_dataset: Dataset = None):
        """Train the NER model"""
        training_args = TrainingArguments(
            output_dir="./models/ner",
            num_train_epochs=3,
            per_device_train_batch_size=8,
            per_device_eval_batch_size=8,
            warmup_steps=500,
            weight_decay=0.01,
            logging_dir="./logs",
            logging_steps=10,
            evaluation_strategy="epoch" if eval_dataset else "no",
        )

        data_collator = DataCollatorForTokenClassification(self.ner_tokenizer)

        trainer = Trainer(
            model=self.ner_model,
            args=training_args,
            train_dataset=train_dataset,
            eval_dataset=eval_dataset,
            data_collator=data_collator,
        )

        trainer.train()
        trainer.save_model()

    def train_sentence_transformer(self, original_texts: List[str], clean_texts: List[str]):
        """Train the sentence transformer model"""
        train_examples = []
        for orig, clean in zip(original_texts, clean_texts):
            train_examples.append(InputExample(texts=[orig, clean], label=1.0))

        train_dataloader = DataLoader(train_examples, shuffle=True, batch_size=16)
        train_loss = losses.CosineSimilarityLoss(self.sentence_transformer)

        self.sentence_transformer.fit(
            train_objectives=[(train_dataloader, train_loss)],
            epochs=3,
            warmup_steps=100,
            show_progress_bar=True
        )

        self.sentence_transformer.save("./models/sentence_transformer")

    async def train_models(self, training_data: List[Dict[str, bytes]]):
        """Train all models using the provided training data"""
        # Process DOCX files
        processed_data = []
        
        for item in training_data:
            if "redline" in item:
                # If we have a redline version, extract changes
                changes = self.extract_changes_from_redline(item["redline"])
                processed_data.extend(changes)
            else:
                # If we have original and clean versions
                original_text = self.extract_text_from_docx(item["original"])
                clean_text = self.extract_text_from_docx(item["clean"])
                processed_data.append({
                    "original": original_text,
                    "redline": clean_text,  # Use clean text as redline for training
                    "clean": clean_text
                })

        # Extract texts from processed data
        original_texts = [item["original"] for item in processed_data]
        redline_texts = [item["redline"] for item in processed_data]
        clean_texts = [item["clean"] for item in processed_data]

        # Train classifier
        classifier_dataset = self.prepare_classification_dataset(original_texts, redline_texts)
        self.train_classifier(classifier_dataset)

        # Train NER (assuming we have labeled data for clause boundaries)
        # This would need to be implemented based on your specific needs
        # ner_dataset = self.prepare_ner_dataset(texts, labels)
        # self.train_ner(ner_dataset)

        # Train sentence transformer
        self.train_sentence_transformer(original_texts, clean_texts)

        return {
            "status": "success",
            "message": "Models trained successfully",
            "models_saved": [
                "./models/classifier",
                "./models/sentence_transformer"
            ],
            "training_samples": len(processed_data)
        } 